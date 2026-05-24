# backend/system_tools/reader.py
import base64
import asyncio
import mimetypes
from pathlib import Path
from typing import Optional, Union, Dict, Any
from dataclasses import dataclass, asdict

from config_loader import config
from backend.utils.validators import validate_path
import backend


# ============ 自定义异常 ============
class FileReadError(Exception):
    """
    文件读取过程中出现的错误，将被 MCP 框架转为 isError 响应。
    """
    pass


@dataclass
class FileReadResult:
    """文件读取结果"""
    content: str
    format: str          # 原始格式: text, markdown, json, csv, binary, image 等
    mime_type: str
    metadata: Dict[str, Any]  # 额外元信息（如 sheet 名、页数、图片尺寸等）


# ============ 可选依赖的懒加载 ============
def _get_markitdown():
    try:
        from markitdown import MarkItDown
        return MarkItDown()
    except ImportError:
        return None

def _get_pandas():
    try:
        import pandas as pd
        return pd
    except ImportError:
        return None

def _get_docx():
    try:
        import docx
        return docx
    except ImportError:
        return None

def _get_pypdf():
    try:
        import PyPDF2
        return PyPDF2
    except ImportError:
        return None

def _get_chardet():
    try:
        import chardet
        return chardet
    except ImportError:
        return None


# ============ 通用工具函数 ============
def df_to_markdown(df) -> str:
    """将 DataFrame 转为 Markdown 表格（不依赖 tabulate，支持清理特殊字符）"""
    cols = [str(c).replace('\n', ' ').replace('|', '\\|') for c in df.columns]
    lines = []
    
    # 表头与分隔行
    lines.append('| ' + ' | '.join(cols) + ' |')
    lines.append('|' + '|'.join(['---'] * len(cols)) + '|')
    
    # 数据行
    for _, row in df.iterrows():
        # 清理单元格中的换行符和管道符，防止打乱 Markdown 表格布局
        row_strs = [str(v).replace('\n', '<br>').replace('|', '\\|') for v in row.values]
        lines.append('| ' + ' | '.join(row_strs) + ' |')
        
    return '\n'.join(lines)


# ============ 格式处理器 ============
class FormatHandler:
    @staticmethod
    def can_handle(path: Path, mime_type: str) -> bool:
        raise NotImplementedError
    
    @staticmethod
    def read(path: Path, encoding: Optional[str] = None, **kwargs) -> FileReadResult:
        raise NotImplementedError


class TextHandler(FormatHandler):
    """纯文本处理器"""
    TEXT_EXTENSIONS = {
        '.txt', '.md', '.markdown', '.rst', '.py', '.js', '.ts', '.jsx', '.tsx', '.vue',
        '.json', '.yaml', '.yml', '.xml', '.html', '.htm', '.css', '.scss', '.less',
        '.sh', '.bash', '.zsh', '.fish', '.ps1', '.bat', '.cmd',
        '.sql', '.c', '.cpp', '.h', '.hpp', '.java', '.go', '.rs', '.rb', '.php',
        '.swift', '.kt', '.scala', '.r', '.m', '.mm', '.pl', '.lua', '.vim',
        '.dockerfile', '.gitignore', '.env', '.ini', '.cfg', '.conf', '.properties',
        '.log', '.svg',
    }
    
    @classmethod
    def can_handle(cls, path: Path, mime_type: str) -> bool:
        if path.suffix.lower() in cls.TEXT_EXTENSIONS:
            return True
        if mime_type and mime_type.startswith(('text/', 'application/json', 'application/xml')):
            return True
        return False
    
    @classmethod
    def read(cls, path: Path, encoding: Optional[str] = None, **kwargs) -> FileReadResult:
        preferred_enc = encoding or 'utf-8'
        actual_enc = preferred_enc
        
        try:
            content = path.read_text(encoding=preferred_enc)
        except UnicodeDecodeError:
            # 统一读取一次字节到内存，避免多次磁盘 I/O
            raw = path.read_bytes()
            chardet = _get_chardet()
            
            if chardet:
                detected = chardet.detect(raw)
                detected_enc = detected.get('encoding') or 'utf-8'
                try:
                    content = raw.decode(detected_enc)
                    actual_enc = detected_enc
                except UnicodeDecodeError:
                    content = raw.decode(detected_enc, errors='replace')
                    actual_enc = detected_enc
            else:
                for try_enc in ['utf-8', 'gbk', 'gb18030', 'latin-1', 'cp1252']:
                    try:
                        content = raw.decode(try_enc)
                        actual_enc = try_enc
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    content = raw.decode('utf-8', errors='replace')
                    actual_enc = 'utf-8'
        
        return FileReadResult(
            content=content,
            format='text',
            mime_type=kwargs.get('mime_type', 'text/plain'),
            metadata={'encoding': actual_enc}
        )


class MarkdownHandler(FormatHandler):
    """通用文档转 Markdown 处理器"""
    SUPPORTED_EXTS = {'.pdf', '.docx', '.doc', '.pptx', '.ppt', '.xlsx', '.xls'}
    
    @classmethod
    def can_handle(cls, path: Path, mime_type: str) -> bool:
        return path.suffix.lower() in cls.SUPPORTED_EXTS
    
    @classmethod
    def read(cls, path: Path, encoding: Optional[str] = None, **kwargs) -> FileReadResult:
        md = _get_markitdown()
        if md is None:
            raise ImportError(
                f"读取 {path.suffix} 文件需要安装 markitdown: pip install 'markitdown[all]'"
            )
        
        try:
            result = md.convert(str(path))
        except ImportError:
            raise
        except Exception as e:
            raise ImportError(f"markitdown 转换失败: {e}") from e
        
        return FileReadResult(
            content=result.text_content,
            format='markdown',
            mime_type=kwargs.get('mime_type', 'text/markdown'),
            metadata={'title': result.title, 'source': str(path)}
        )


class CSVHandler(FormatHandler):
    """CSV/TSV 处理器"""
    @classmethod
    def can_handle(cls, path: Path, mime_type: str) -> bool:
        return path.suffix.lower() in {'.csv', '.tsv'}
    
    @classmethod
    def read(cls, path: Path, encoding: Optional[str] = None, **kwargs) -> FileReadResult:
        pd = _get_pandas()
        if pd is None:
            return TextHandler.read(path, encoding, **kwargs)
        
        enc = encoding or 'utf-8'
        sep = '\t' if path.suffix.lower() == '.tsv' else ','
        
        try:
            df = pd.read_csv(path, encoding=enc, sep=sep)
        except UnicodeDecodeError:
            for try_enc in ['gbk', 'gb18030', 'latin-1']:
                try:
                    df = pd.read_csv(path, encoding=try_enc, sep=sep)
                    enc = try_enc
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise FileReadError(f"无法解码 CSV 文件: {path}")
        
        return FileReadResult(
            content=df_to_markdown(df),
            format='markdown',
            mime_type='text/markdown',
            metadata={
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': list(df.columns),
                'encoding': enc,
            }
        )


class ExcelHandler(FormatHandler):
    """Excel 处理器，支持多 sheet"""
    @classmethod
    def can_handle(cls, path: Path, mime_type: str) -> bool:
        return path.suffix.lower() in {'.xlsx', '.xls', '.xlsm', '.xlsb'}
    
    @classmethod
    def read(cls, path: Path, encoding: Optional[str] = None, 
             sheet_name: Optional[Union[str, int]] = None, **kwargs) -> FileReadResult:
        pd = _get_pandas()
        if pd is None:
            raise ImportError("读取 Excel 需要安装 pandas: pip install pandas openpyxl")
        
        ext = path.suffix.lower()
        engine = 'xlrd' if ext == '.xls' else 'openpyxl'
        
        try:
            if sheet_name is not None:
                df = pd.read_excel(path, sheet_name=sheet_name, engine=engine)
                return FileReadResult(
                    content=df_to_markdown(df),
                    format='markdown',
                    mime_type='text/markdown',
                    metadata={
                        'sheets': [sheet_name],
                        'rows': len(df),
                        'columns': len(df.columns),
                    }
                )
            else:
                with pd.ExcelFile(path, engine=engine) as xl:
                    sheets_info = xl.sheet_names
                
                if len(sheets_info) == 1:
                    df = pd.read_excel(path, sheet_name=0, engine=engine)
                    content = df_to_markdown(df)
                else:
                    parts = []
                    for name in sheets_info:
                        df_sheet = pd.read_excel(path, sheet_name=name, engine=engine)
                        parts.append(f"## Sheet: {name}\n\n{df_to_markdown(df_sheet)}\n\n")
                    content = ''.join(parts)
                    
                return FileReadResult(
                    content=content,
                    format='markdown',
                    mime_type='text/markdown',
                    metadata={'sheets': sheets_info}
                )
                
        except ImportError:
            raise ImportError(f"读取 {ext} 需要安装对应引擎：pip install {engine}") from None
        except Exception as e:
            md = _get_markitdown()
            if md:
                try:
                    result = md.convert(str(path))
                    return FileReadResult(
                        content=result.text_content,
                        format='markdown',
                        mime_type='text/markdown',
                        metadata={'fallback': 'markitdown'}
                    )
                except Exception:
                    pass
            raise FileReadError(f"Excel 读取失败: {e}") from e


class DocxHandler(FormatHandler):
    """Word 文档处理器"""
    @classmethod
    def can_handle(cls, path: Path, mime_type: str) -> bool:
        return path.suffix.lower() == '.docx'
    
    @classmethod
    def read(cls, path: Path, encoding: Optional[str] = None, **kwargs) -> FileReadResult:
        docx = _get_docx()
        if docx is None:
            raise ImportError("读取 DOCX 需要安装 python-docx: pip install python-docx")
        
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = '\n\n'.join(paragraphs)
        
        tables_md = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                # 修复: 清理换行符，防止 Markdown 表格断裂
                cells = [cell.text.replace('\n', '<br>').replace('|', '\\|') for cell in row.cells]
                rows.append('| ' + ' | '.join(cells) + ' |')
            if rows:
                header_sep = '|' + '|'.join(['---'] * len(table.rows[0].cells)) + '|'
                rows.insert(1, header_sep)
                tables_md.append('\n'.join(rows))
        
        if tables_md:
            content += '\n\n' + '\n\n'.join(tables_md)
        
        return FileReadResult(
            content=content,
            format='markdown',
            mime_type='text/markdown',
            metadata={'paragraphs': len(paragraphs), 'tables': len(doc.tables)}
        )


class PDFHandler(FormatHandler):
    """PDF 处理器"""
    @classmethod
    def can_handle(cls, path: Path, mime_type: str) -> bool:
        return path.suffix.lower() == '.pdf'
    
    @classmethod
    def read(cls, path: Path, encoding: Optional[str] = None, **kwargs) -> FileReadResult:
        PyPDF2 = _get_pypdf()
        if PyPDF2 is None:
            raise ImportError("读取 PDF 需要安装 PyPDF2: pip install PyPDF2")
        
        text_parts = []
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"## Page {i + 1}\n\n{text}")
                except Exception:
                    text_parts.append(f"## Page {i + 1}\n\n[无法提取文本]")
        
        return FileReadResult(
            content='\n\n'.join(text_parts),
            format='markdown',
            mime_type='text/markdown',
            metadata={'pages': len(reader.pages)}
        )


class ImageHandler(FormatHandler):
    """图片处理器"""
    IMAGE_EXTS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.tiff', '.ico'}
    
    @classmethod
    def can_handle(cls, path: Path, mime_type: str) -> bool:
        if path.suffix.lower() in cls.IMAGE_EXTS:
            return True
        if mime_type and mime_type.startswith('image/'):
            return True
        return False
    
    @classmethod
    def read(cls, path: Path, encoding: Optional[str] = None, **kwargs) -> FileReadResult:
        mime_type = kwargs.get('mime_type', 'image/png')
        
        # 优化: 仅在需要 base64 时读取，通过 kwargs['return_base64'] 控制
        if kwargs.get('return_base64', True):
            data = path.read_bytes()
            b64 = base64.b64encode(data).decode('ascii')
            content = f"data:{mime_type};base64,{b64}"
            size = len(data)
        else:
            content = f"![{path.name}]({path.name})"
            size = path.stat().st_size
            
        return FileReadResult(
            content=content,
            format='image',
            mime_type=mime_type,
            metadata={'size_bytes': size}
        )


class BinaryHandler(FormatHandler):
    """二进制文件回退处理器"""
    @classmethod
    def can_handle(cls, path: Path, mime_type: str) -> bool:
        return True 
    
    @classmethod
    def read(cls, path: Path, encoding: Optional[str] = None, **kwargs) -> FileReadResult:
        stat = path.stat()
        mime_type = kwargs.get('mime_type', 'application/octet-stream')
        
        if stat.st_size <= 1024:
            data = path.read_bytes()
            hex_content = data.hex(' ')
            content = f"Binary file ({mime_type})\nSize: {stat.st_size} bytes\nHex: {hex_content}"
        else:
            content = f"Binary file ({mime_type})\nSize: {stat.st_size} bytes\n[内容太大，无法以文本形式显示]"
        
        return FileReadResult(
            content=content,
            format='binary',
            mime_type=mime_type,
            metadata={'size_bytes': stat.st_size}
        )


# ============ 处理器优先级列表 ============
HANDLERS = [
    ImageHandler,
    CSVHandler,
    MarkdownHandler,
    ExcelHandler,
    DocxHandler,
    PDFHandler,
    TextHandler,
    BinaryHandler,
]


# ============ 主函数 ============
async def file_read(
    path: str,
    sheet_name: Optional[Union[str, int]] = None,
    encoding: str = "UTF-8",
    max_size_mb: int = 10,
    return_image_base64: bool = False
) -> dict[str, Any]:
    """
    读取指定路径的文件内容
    """
    paths = [config.uploads_dir, backend.workspace_path]
    allowed_dirs = [Path(p).resolve() for p in paths]

    max_bytes = max_size_mb * 1024 * 1024

    path_obj = Path(path)
    if not path_obj.is_absolute():
        path_obj = Path.cwd() / path_obj
    
    try:
        safe_path = validate_path(str(path_obj), allowed_dirs)
    except ValueError as e:
        raise FileReadError(f"路径校验失败：{e}") from e

    if not safe_path.exists():
        raise FileReadError(f"文件不存在：{safe_path}")
    if not safe_path.is_file():
        raise FileReadError(f"路径不是文件：{safe_path}")

    try:
        file_size = safe_path.stat().st_size
    except OSError as e:
        raise FileReadError(f"无法获取文件信息：{e}") from e

    if file_size > max_bytes:
        raise FileReadError(
            f"文件过大（{file_size} 字节），超过限制 {max_bytes} 字节。"
        )

    mime_type, _ = mimetypes.guess_type(str(safe_path))
    mime_type = mime_type or 'application/octet-stream'

    loop = asyncio.get_running_loop()

    def _sync_read():
        last_import_error = None
        for handler in HANDLERS:
            if not handler.can_handle(safe_path, mime_type):
                continue
            try:
                result = handler.read(
                    safe_path,
                    encoding=encoding,
                    sheet_name=sheet_name,
                    mime_type=mime_type,
                    return_base64=return_image_base64
                )

                # ========= 核心修改点 =========
                # 将 dataclass 转换为字典
                result_dict = asdict(result)
                
                # 序列化为 JSON 字符串返回 (ensure_ascii=False 保证中文正常显示)
                return result_dict
                # ===============================
                
            except ImportError as e:
                last_import_error = e
                continue
            except Exception as e:
                if handler is BinaryHandler:
                    raise FileReadError(f"读取文件时发生致命错误：{e}") from e
                continue
                
        if last_import_error:
            raise FileReadError(
                f"缺少必要的依赖，无法读取该文件: {last_import_error}"
            ) from last_import_error
        raise FileReadError("无法读取文件：没有合适的处理器")

    return await loop.run_in_executor(None, _sync_read)